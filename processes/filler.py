from docx import Document
from docx.shared import RGBColor

from datetime import datetime


def fill(row, cell: int, header_data: dict[str, str], data: dict[str, str]) -> None:
    curr_cell = row.cells[cell]
    text = curr_cell.text.strip()

    if header_data.get(text) is not None:
        curr_cell.text = header_data[text]

    if data.get(text) is not None:
        next_cell = row.cells[cell + 1]
        next_cell.text = data[text]

    return


def fill_doc(
    file_name: str,
    header_data: dict[str, str],
    data: dict[str, str],
    date_time: datetime,
) -> list[str]:
    doc = Document(file_name)

    for table in doc.tables:
        for row in table.rows:
            for c in range(len(row.cells)):
                fill(row, c, header_data, data)

    timestamp = date_time.strftime("%Y-%m-%d_%H-%M")
    doc.paragraphs[len(doc.paragraphs) - 1].text = f"Edited by script on {timestamp}"
    run = doc.paragraphs[len(doc.paragraphs) - 1].runs[0]
    run.font.color.rgb = RGBColor(255, 255, 255)
    doc.save(file_name)
