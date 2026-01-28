from docx import Document
from docx.shared import RGBColor
from Person import Person
from Project import Project

INCOMPLETE_DATA = "Ostatné údaje nie sú dostupné v žiadosti ani v PD."

def fill(row, cell: int, header_data: dict[str, str], data: dict[str, str]) -> None:
    curr_cell = row.cells[cell]
    text = curr_cell.text.strip()

    if header_data.get(text) is not None:
        curr_cell.text = header_data[text]
    
    if data.get(text) is not None:
        next_cell = row.cells[cell + 1]
        next_cell.text = data[text]
    
    return


def prepare_applicants(applicants: list[Person],
                       header_data: dict[str, str],
                       data: dict[str, str]) -> None:
    for applicant in applicants:
        header_data["Meno Priezvisko"] = \
            header_data.get("Meno Priezvisko", "") + f"{applicant.get_full_name()}"
        
        header_data["Ulica číslo"] = \
            header_data.get("Ulica číslo", "") + f"{applicant.address.street} {applicant.address.building_number}"
        
        header_data["PSČ Mesto"] = \
            header_data.get("PSČ Mesto", "") + f"{applicant.address.postal_code} {applicant.address.city} \n"
        
        data["Žiadateľ"] = \
            data.get("Žiadateľ", "") + f"{applicant.to_string()} \n"
        
        if not applicant.is_complete():
             data["Žiadateľ"] += f"\n{INCOMPLETE_DATA}\n"
    return


def prepare_project_owners(project_owners: list[Person],
                           data: dict[str, str]) -> None:
    for project_owner in project_owners:
        data["Stavebník"] = \
            data.get("Stavebník", "") + f"{project_owner.to_string()} \n"
        
        if not project_owner.is_complete():
             data["Stavebník"] += f"\n{INCOMPLETE_DATA}\n"


def prepare_project(project: Project, data: dict[str, str]) -> None:
    data["ID stavby \n(ak bolo pridelené Portálom výstavby)"] = project.id
    data["Názov stavby"] = project.title
    data["Identifikačné údaje stavby"] = project.format_parcels()
    data["Členenie stavby"] = project.format_facilities()
    return


def prepare_data(applicants: list[Person],
                 project_owners: list[Person],
                 project: Project) -> tuple[dict[str, str], dict[str, str]]:
    header_data: dict[str, str] = {}
    data: dict[str, str] = {}

    prepare_applicants(applicants, header_data, data)
    prepare_project_owners(project_owners, data)
    prepare_project(project, data)

    return header_data, data

def fill_doc(file_name: str, applicants: list[Person], project_owners: list[Person],
             project: Project, timestamp: str) -> list[str]:
    doc = Document(file_name)
    header_data, data = prepare_data(applicants, project_owners, project)

    for table in doc.tables:
        for row in table.rows:
            for c in range(len(row.cells)):
                fill(row, c, header_data, data)
    
    doc.paragraphs[len(doc.paragraphs) - 1].text = f"Edited by script on {timestamp}"
    run = doc.paragraphs[len(doc.paragraphs) - 1].runs[0]
    run.font.color.rgb = RGBColor(255, 255, 255)
    doc.save(file_name)


